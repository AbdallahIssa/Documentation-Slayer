#!/usr/bin/env python3
import re
import json
import sys
import argparse
import os
import platform
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill
import subprocess
import threading

# PyQt6 imports
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                              QTabWidget, QPushButton, QLabel, QLineEdit, QCheckBox,
                              QFileDialog, QMessageBox, QTableWidget, QTableWidgetItem,
                              QProgressDialog, QDialog, QDialogButtonBox, QGroupBox,
                              QGridLayout, QHeaderView, QStyle, QComboBox)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt6.QtGui import QIcon, QFont, QColor

# Try to import tqdm for CLI progress bars (optional)
try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    tqdm = None


"""
I'm using This command to get the executable for the parser script:
pyinstaller --onefile --windowed --name "Doc-Slayer" --icon "DocSlayerLogo.ico" --add-data "DocSlayerLogo.ico;." --add-data "CodeSmasher.exe;." --add-data "qt_gui_modern.py;." --hidden-import PyQt6 parser.py
"""

# Global cancellation flag
class CancellationToken:
    """Thread-safe cancellation token for long-running operations"""
    def __init__(self):
        self.cancelled = False
        self._lock = threading.Lock()

    def cancel(self):
        with self._lock:
            self.cancelled = True

    def is_cancelled(self):
        with self._lock:
            return self.cancelled

    def reset(self):
        with self._lock:
            self.cancelled = False


class ParserThread(QThread):
    """Worker thread for parsing operations"""
    finished = pyqtSignal(bool, str, list, list, list)  # success, error, functions, macros, variables
    progress = pyqtSignal(str)  # progress text

    def __init__(self, file_path, cancel_token):
        super().__init__()
        self.file_path = file_path
        self.cancel_token = cancel_token

    def run(self):
        try:
            self.progress.emit("Reading file...")
            with open(self.file_path, encoding="utf-8") as f:
                src = f.read()

            self.progress.emit(f"Parsing file ({len(src):,} bytes)...")
            functions, macros, variables = parse_file(src, self.cancel_token)

            if self.cancel_token.is_cancelled():
                self.finished.emit(False, "Operation cancelled", [], [], [])
                return

            self.finished.emit(True, "", functions, macros, variables)
        except Exception as e:
            self.finished.emit(False, str(e), [], [], [])


class PasswordManager:
    """Manages password authentication with session persistence"""
    def __init__(self):
        self.is_authenticated = False
        self.attempts = 0
        self.max_attempts = 3
    
    def reset(self):
        """Reset authentication state"""
        self.is_authenticated = False
        self.attempts = 0


# Global password manager instance
password_manager = PasswordManager()


class PasswordDialog(QDialog):
    """Password dialog for Activity Diagram tab access"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Access Activity Diagram")
        self.setModal(True)
        self.setFixedSize(400, 200)

        layout = QVBoxLayout()

        # Title
        title = QLabel("Enter password to access Activity Diagram:")
        title.setFont(QFont("Arial", 10))
        layout.addWidget(title)

        # Attempts remaining
        attempts_left = password_manager.max_attempts - password_manager.attempts
        self.attempts_label = QLabel(f"Attempts remaining: {attempts_left}")
        layout.addWidget(self.attempts_label)

        # Password input
        self.password_input = QLineEdit()
        self.password_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.password_input.setPlaceholderText("Enter password...")
        self.password_input.returnPressed.connect(self.check_password)
        layout.addWidget(self.password_input)

        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(self.check_password)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

        self.setLayout(layout)
        self.access_granted = False

    def check_password(self):
        password_manager.attempts += 1

        if self.password_input.text() == "Vehiclevo@1234":
            password_manager.is_authenticated = True
            self.access_granted = True
            self.accept()
        else:
            if password_manager.attempts >= password_manager.max_attempts:
                QMessageBox.critical(self, "Access Denied",
                                   "Maximum password attempts exceeded. Application will close.")
                self.reject()
                QApplication.quit()
            else:
                remaining = password_manager.max_attempts - password_manager.attempts
                QMessageBox.warning(self, "Invalid Password",
                                  f"Wrong password! {remaining} attempt(s) remaining.")
                self.password_input.clear()
                self.attempts_label.setText(f"Attempts remaining: {remaining}")


def ask_password(parent_window=None):
    """Ask for password to access Activity Diagram tab with 3 attempts"""
    if password_manager.is_authenticated:
        return True

    if password_manager.attempts >= password_manager.max_attempts:
        QMessageBox.critical(parent_window, "Access Denied",
                           "Maximum password attempts exceeded. Application will close.")
        QApplication.quit()
        return False

    dialog = PasswordDialog(parent_window)
    dialog.exec()
    return dialog.access_granted


def open_file(path: str):
    """Open a file with the default OS application."""
    system = platform.system()
    if system == "Windows":
        os.startfile(path)
    elif system == "Darwin":
        os.system(f'open "{path}"')
    else:
        os.system(f'xdg-open "{path}"')

def write_excel(file_path: str, functions: list[dict], macros: list[dict], variables: list[dict], 
                sel_function_fields: list[str], sel_macro_fields: list[str], sel_variable_fields: list[str]):
    """
    Write an .xlsx with three sheets: Functions, Macros, and Variables.
    Only the selected fields columns are written for each sheet.
    Headers are styled bold+red font on yellow fill.
    """
    FUNCTION_HEADERS = [
      'Line Number', 'Name', 'Syntax', 'Return Value', 'In-Parameters', 'Out-Parameters',
      'Function Type', 'Description', 'Sync/Async', 'Reentrancy',
      'Triggers', 'Inputs', 'Outputs',
      'Invoked Operations', 'Used Data Types'
    ]

    MACRO_HEADERS = [
        'Line Number', 'Name', 'Value'
    ]

    VARIABLE_HEADERS = [
        'Line Number', 'Name', 'Data Type', 'Initial Value', 'Scope'
    ]

    path = Path(file_path)
    if path.exists():
        wb = load_workbook(str(path))
        # Remove all existing sheets
        for name in list(wb.sheetnames):
            wb.remove(wb[name])
    else:
        wb = Workbook()
        # Remove default sheet
        wb.remove(wb.active)

    yellow = PatternFill(fill_type="solid", fgColor="FFFFFF00")
    red = Font(bold=True, color="FFFF0000")

    # Functions sheet
    if functions:
        ws_functions = wb.create_sheet(title="Runnables and static functions")
        function_headers = [h for h in FUNCTION_HEADERS if h in sel_function_fields]
        ws_functions.append(function_headers)
        
        for cell in ws_functions[1]:
            cell.fill = yellow
            cell.font = red

        for r in functions:
            row = []
            for h in function_headers:
                if h == 'Line Number':
                    row.append(r.get('lineNumber', ''))
                elif h == 'Name':
                    row.append(r['name'])
                elif h == 'Syntax':
                    row.append(r['syntax'])
                elif h == 'Return Value':
                    row.append(r['ret'])
                elif h == 'In-Parameters':
                    row.append(", ".join(r['inParams']))
                elif h == 'Out-Parameters':
                    row.append(", ".join(r['outParams']))
                elif h == 'Function Type':
                    row.append(r['fnType'])
                elif h == 'Description':
                    row.append(r.get('description', ''))
                elif h == 'Sync/Async':
                    row.append(r['Sync_Async'])
                elif h == 'Reentrancy':
                    row.append(r['Reentrancy'])
                elif h == 'Triggers':
                    row.append(r['trigger'])
                elif h == 'Inputs':
                    row.append(", ".join(r['inputs']))
                elif h == 'Outputs':
                    row.append(", ".join(r['outputs']))
                elif h == 'Invoked Operations':
                    row.append(", ".join(r['invoked']))
                elif h == 'Used Data Types':
                    row.append(", ".join(r['used']))
            ws_functions.append(row)

        # Auto-adjust column widths
        for col in ws_functions.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                val = str(cell.value or "")
                max_length = max(max_length, len(val))
            ws_functions.column_dimensions[col_letter].width = max_length + 2

    # Macros sheet
    if macros:
        ws_macros = wb.create_sheet(title="Macros")
        macro_headers = [h for h in MACRO_HEADERS if h in sel_macro_fields]
        ws_macros.append(macro_headers)
        
        for cell in ws_macros[1]:
            cell.fill = yellow
            cell.font = red

        for r in macros:
            row = []
            for h in macro_headers:
                if h == 'Line Number':
                    row.append(r.get('lineNumber', ''))
                elif h == 'Name':
                    row.append(r['name'])
                elif h == 'Value':
                    row.append(r['value'])
            ws_macros.append(row)

        # Auto-adjust column widths
        for col in ws_macros.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                val = str(cell.value or "")
                max_length = max(max_length, len(val))
            ws_macros.column_dimensions[col_letter].width = max_length + 2

    # Variables sheet
    if variables:
        ws_variables = wb.create_sheet(title="Variables")
        variable_headers = [h for h in VARIABLE_HEADERS if h in sel_variable_fields]
        ws_variables.append(variable_headers)
        
        for cell in ws_variables[1]:
            cell.fill = yellow
            cell.font = red

        for r in variables:
            row = []
            for h in variable_headers:
                if h == 'Line Number':
                    row.append(r.get('lineNumber', ''))
                elif h == 'Name':
                    row.append(r['name'])
                elif h == 'Data Type':
                    row.append(r['dataType'])
                elif h == 'Initial Value':
                    row.append(r['initialValue'])
                elif h == 'Scope':
                    row.append(r['scope'])
            ws_variables.append(row)

        # Auto-adjust column widths
        for col in ws_variables.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                val = str(cell.value or "")
                max_length = max(max_length, len(val))
            ws_variables.column_dimensions[col_letter].width = max_length + 2

    wb.save(str(path))

def write_markdown(file_path: str, functions: list[dict], macros: list[dict], variables: list[dict],
                   sel_function_fields: list[str], sel_macro_fields: list[str], sel_variable_fields: list[str]):
    """
    Write a Markdown file with tables for functions, macros, and variables.
    Only the selected fields are included in each table.
    """
    FUNCTION_FIELD_GETTERS = {
      'Line Number':      lambda r: str(r.get('lineNumber', '')),
      'Name':             lambda r: r['name'],
      'Syntax':           lambda r: f"`{r['syntax']}`",
      'Sync/Async':       lambda r: f"`{r['Sync_Async']}`",
      'Reentrancy':       lambda r: f"`{r['Reentrancy']}`",
      'Return Value':     lambda r: f"`{r['ret']}`",
      'In-Parameters':    lambda r: ", ".join(r['inParams']),
      'Out-Parameters':   lambda r: ", ".join(r['outParams']),
      'Function Type':    lambda r: r['fnType'],
      'Description':      lambda r: r.get('description', ''),
      'Triggers':         lambda r: r['trigger'],
      'Inputs':           lambda r: ", ".join(r['inputs']),
      'Outputs':          lambda r: ", ".join(r['outputs']),
      'Invoked Operations': lambda r: ", ".join(r['invoked']),
      'Used Data Types':    lambda r: ", ".join(r['used']),
    }

    MACRO_FIELD_GETTERS = {
        'Line Number': lambda r: str(r.get('lineNumber', '')),
        'Name':  lambda r: r['name'],
        'Value': lambda r: f"`{r['value']}`",
    }

    VARIABLE_FIELD_GETTERS = {
        'Line Number':   lambda r: str(r.get('lineNumber', '')),
        'Name':         lambda r: r['name'],
        'Data Type':    lambda r: f"`{r['dataType']}`",
        'Initial Value': lambda r: f"`{r['initialValue']}`",
        'Scope':        lambda r: r['scope'],
    }

    lines = []

    # Functions section
    if functions:
        lines.append("# Functions")
        lines.append(f"**Selected fields:** {', '.join(sel_function_fields)}")
        lines.append("")

        for r in functions:
            lines.append(f"## {r['name']}")
            lines.append("")
            lines.append("| Field | Value |")
            lines.append("|-------|-------|")
            for label in sel_function_fields:
                getter = FUNCTION_FIELD_GETTERS.get(label)
                if getter:
                    value = getter(r)
                    lines.append(f"| {label} | {value} |")
            lines.append("")

    # Macros section
    if macros:
        lines.append("# Macros")
        lines.append(f"**Selected fields:** {', '.join(sel_macro_fields)}")
        lines.append("")

        for r in macros:
            lines.append(f"## {r['name']}")
            lines.append("")
            lines.append("| Field | Value |")
            lines.append("|-------|-------|")
            for label in sel_macro_fields:
                getter = MACRO_FIELD_GETTERS.get(label)
                if getter:
                    value = getter(r)
                    lines.append(f"| {label} | {value} |")
            lines.append("")

    # Variables section
    if variables:
        lines.append("# Variables")
        lines.append(f"**Selected fields:** {', '.join(sel_variable_fields)}")
        lines.append("")

        for r in variables:
            lines.append(f"## {r['name']}")
            lines.append("")
            lines.append("| Field | Value |")
            lines.append("|-------|-------|")
            for label in sel_variable_fields:
                getter = VARIABLE_FIELD_GETTERS.get(label)
                if getter:
                    value = getter(r)
                    lines.append(f"| {label} | {value} |")
            lines.append("")

    Path(file_path).write_text("\n".join(lines), encoding="utf-8")

def write_docx(source_path: str, swcName: str, functions: list[dict], macros: list[dict], variables: list[dict],
               sel_function_fields: list[str], sel_macro_fields: list[str], sel_variable_fields: list[str]):
    """
    Generate a Word .docx per AUTOSAR template.
    Includes functions, macros, and variables with only selected fields.
    """
    def shade_cell(cell, rgb="9D9D9D"):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), rgb)
        tcPr.append(shd)

    doc = Document()

    # Functions section
    if functions:
        doc.add_heading('Functions', level=1)
        
        for r in functions:
            doc.add_paragraph(f"[{r['name']}]", style='Heading 2')
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = 'Table Grid'

            def add_row(label, value):
                c0, c1 = tbl.add_row().cells
                c0.text = label
                c1.text = value or ""
                shade_cell(c0)

            if 'Line Number' in sel_function_fields:
                add_row("Line Number", str(r.get('lineNumber', '')))
            if 'Name' in sel_function_fields:
                add_row("Service Name", r['name'])
            if 'Syntax' in sel_function_fields:
                add_row("Syntax", r['syntax'])
            if 'Sync/Async' in sel_function_fields:
                add_row("Sync/Async", r['Sync_Async'])
            if 'Reentrancy' in sel_function_fields:
                add_row("Reentrancy", r['Reentrancy'])
            if 'In-Parameters' in sel_function_fields:
                add_row("Parameters (in)", ", ".join(r['inParams']))
            if 'Out-Parameters' in sel_function_fields:
                add_row("Parameters (out)", ", ".join(r['outParams']))
            if 'Function Type' in sel_function_fields:
                add_row("Function Type", r['fnType'])
            if 'Description' in sel_function_fields:
                add_row("Description", r.get('description', ''))
            if 'Triggers' in sel_function_fields:
                add_row("Triggers", r['trigger'])
            if 'Inputs' in sel_function_fields:
                add_row("Inputs", ", ".join(r['inputs']))
            if 'Outputs' in sel_function_fields:
                add_row("Outputs", ", ".join(r['outputs']))
            if 'Invoked Operations' in sel_function_fields:
                add_row("Invoked Operations", ", ".join(r['invoked']))
            if 'Used Data Types' in sel_function_fields:
                add_row("Used Data Types", ", ".join(r['used']))

            doc.add_page_break()

    # Macros section
    if macros:
        doc.add_heading('Macros', level=1)
        
        for r in macros:
            doc.add_paragraph(f"[{r['name']}]", style='Heading 2')
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = 'Table Grid'

            def add_row(label, value):
                c0, c1 = tbl.add_row().cells
                c0.text = label
                c1.text = value or ""
                shade_cell(c0)

            if 'Line Number' in sel_macro_fields:
                add_row("Line Number", str(r.get('lineNumber', '')))
            if 'Name' in sel_macro_fields:
                add_row("Macro Name", r['name'])
            if 'Value' in sel_macro_fields:
                add_row("Value", r['value'])

            doc.add_page_break()

    # Variables section
    if variables:
        doc.add_heading('Variables', level=1)
        
        for r in variables:
            doc.add_paragraph(f"[{r['name']}]", style='Heading 2')
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = 'Table Grid'

            def add_row(label, value):
                c0, c1 = tbl.add_row().cells
                c0.text = label
                c1.text = value or ""
                shade_cell(c0)

            if 'Line Number' in sel_variable_fields:
                add_row("Line Number", str(r.get('lineNumber', '')))
            if 'Name' in sel_variable_fields:
                add_row("Variable Name", r['name'])
            if 'Data Type' in sel_variable_fields:
                add_row("Data Type", r['dataType'])
            if 'Initial Value' in sel_variable_fields:
                add_row("Initial Value", r['initialValue'])
            if 'Scope' in sel_variable_fields:
                add_row("Scope", r['scope'])

            doc.add_page_break()

    excel_path = Path(source_path).with_suffix('.xlsx')
    docx_path = excel_path.with_suffix('.docx')
    doc.save(str(docx_path))

def classify_params(body: str, params: list[str], param_types: list[str]) -> dict:
    """
    Precise IN/OUT/INOUT detection for parameters.
    Analyzes pointer usage patterns within the function body.
    """
    result = {}

    for p, ptype in zip(params, param_types):
        esc = re.escape(p)

        # Check if parameter has const qualifier
        if re.search(r'\bconst\b', ptype):
            result[p] = "IN"
            continue

        is_written = False
        is_read = False

        # Pattern 1: Pointer dereference writes - *param = value
        ptr_deref_write = re.search(rf"\*\s*{esc}\s*=", body)
        if ptr_deref_write:
            is_written = True

        # Pattern 2: Array access writes - param[i] = value
        array_write = re.search(rf"\b{esc}\s*\[[^\]]+\]\s*=", body)
        if array_write:
            is_written = True

        # Pattern 3: Structure member writes - param->field = value
        arrow_write = re.search(rf"\b{esc}\s*->\s*\w+\s*=", body)
        if arrow_write:
            is_written = True

        # Pattern 4: Increment/decrement on pointer - ptr++, ++ptr, ptr--, --ptr
        inc_dec = re.search(rf"(?:\+\+|--)\s*{esc}\b|\b{esc}\s*(?:\+\+|--)", body)
        if inc_dec:
            is_written = True

        # Pattern 5: Increment/decrement on dereferenced pointer - (*ptr)++, ++(*ptr)
        deref_inc_dec = re.search(rf"(?:\+\+|--)\s*\(\s*\*\s*{esc}\s*\)|\(\s*\*\s*{esc}\s*\)\s*(?:\+\+|--)", body)
        if deref_inc_dec:
            is_written = True

        # Pattern 6: Assignment to dereferenced pointer in parentheses - (*param) = value
        paren_ptr_write = re.search(rf"\(\s*\*\s*{esc}\s*\)\s*=", body)
        if paren_ptr_write:
            is_written = True

        # Pattern 7: Function calls with param as destination (first parameter typically)
        # memcpy(dest, src, len) - dest is OUT
        memcpy_out = re.search(rf"\b(?:memcpy|strcpy|sprintf|snprintf)\s*\(\s*{esc}\s*,", body)
        if memcpy_out:
            is_written = True

        # Pattern 8: Write APIs - functions with Write/Set in name
        write_api = re.search(rf"\b\w*(?:Write|Set)\w*\s*\([^;]*\b{esc}\b", body)
        if write_api:
            is_written = True

        # === READ PATTERNS ===

        # Pattern 9: Pointer dereference reads - value = *param
        ptr_deref_read = re.search(rf"=\s*\*\s*{esc}\b", body)
        if ptr_deref_read:
            is_read = True

        # Pattern 10: Array access reads - value = param[i]
        array_read = re.search(rf"=\s*{esc}\s*\[[^\]]+\]", body)
        if array_read:
            is_read = True

        # Pattern 11: Structure member reads - value = param->field
        arrow_read = re.search(rf"=\s*{esc}\s*->\s*\w+", body)
        if arrow_read:
            is_read = True

        # Pattern 12: Function calls with param as source (second parameter typically)
        # memcpy(dest, src, len) - src is IN
        memcpy_in = re.search(rf"\b(?:memcpy|strcpy|strcmp|strncmp)\s*\([^,]+,\s*{esc}\s*[,\)]", body)
        if memcpy_in:
            is_read = True

        # Pattern 13: Read APIs - functions with Read/Get in name
        read_api = re.search(rf"\b\w*(?:Read|Get)\w*\s*\([^;]*\b{esc}\b", body)
        if read_api:
            is_read = True

        # Pattern 14: Parameter used in comparison or condition
        condition_use = re.search(rf"(?:if|while|for|switch)\s*\([^)]*\b{esc}\b", body)
        if condition_use:
            is_read = True

        # Pattern 15: Parameter used in expressions (right side of operations)
        expr_use = re.search(rf"[+\-*/%&|^<>]\s*{esc}\b|\b{esc}\s*[+\-*/%&|^<>]", body)
        if expr_use:
            is_read = True

        # Classify based on usage
        if is_written and is_read:
            result[p] = "INOUT"
        elif is_written:
            result[p] = "OUT"
        elif is_read:
            result[p] = "IN"
        else:
            # Parameter not used in body, assume IN
            result[p] = "IN"

    return result

def get_trigger_comment(comments: list, pos: int) -> str:
    trig_rx = re.compile(r"-\s*triggered\s+(?:on|by)\s+([^\n\r]+)", re.IGNORECASE)
    best_end, best_txt = -1, ""
    for end, txt in comments:
        if end <= pos and end > best_end and trig_rx.search(txt):
            best_end, best_txt = end, txt
    return best_txt

def parse_doxygen_comment(comment: str) -> str:
    """
    Parse Doxygen comment and extract description.
    Handles @brief, @details, and combines them.
    """
    if not comment:
        return ""

    # Remove comment delimiters and clean up
    # Handle /** */ style
    text = re.sub(r'/\*\*+!?', '', comment)
    text = re.sub(r'\*/', '', text)
    # Remove leading * from each line
    text = re.sub(r'^\s*\*\s?', '', text, flags=re.MULTILINE)
    # Handle /// style
    text = re.sub(r'^\s*///?!?\s?', '', text, flags=re.MULTILINE)

    # Extract @brief or \brief - stop at next tag (on a newline) or end of text
    brief_match = re.search(r'[@\\]brief\s+(.+?)(?=\s*\n\s*[@\\](?:param|return|details|note|warning|see)\b|\Z)', text, re.DOTALL | re.IGNORECASE)
    brief = brief_match.group(1).strip() if brief_match else ""

    # Extract @details or \details - stop at next tag (on a newline) or end of text
    details_match = re.search(r'[@\\]details\s+(.+?)(?=\s*\n\s*[@\\](?:param|return|brief|note|warning|see)\b|\Z)', text, re.DOTALL | re.IGNORECASE)
    details = details_match.group(1).strip() if details_match else ""

    # If no tags found, try to extract the first paragraph as description
    if not brief and not details:
        # Remove all Doxygen tags and get the first non-empty text before any tag
        no_tags = re.sub(r'[@\\](?:param|return|brief|details|note|warning|see)\b.*', '', text, flags=re.DOTALL)
        lines = [line.strip() for line in no_tags.split('\n') if line.strip()]
        if lines:
            brief = ' '.join(lines)

    # Combine brief and details
    if brief and details:
        # Clean up multiline formatting
        brief = ' '.join(brief.split())
        details = ' '.join(details.split())
        return f"{brief}. {details}"
    elif brief:
        return ' '.join(brief.split())
    elif details:
        return ' '.join(details.split())

    return ""

def get_doxygen_comment(src: str, comments: list, pos: int) -> str:
    """
    Find the closest Doxygen comment before the function position.
    Handles both /** */ and /// style comments.
    """
    # Look for /// style comments first (they're usually directly above the function)
    snippet_before = src[:pos]
    lines = snippet_before.split('\n')

    # Go backwards from the function position
    doxygen_lines = []
    found_function_line = False
    for i in range(len(lines) - 1, -1, -1):
        line = lines[i].strip()

        if not found_function_line and line:
            # This is the function declaration line, skip it
            found_function_line = True
            continue

        if line.startswith('///'):
            doxygen_lines.insert(0, line)
        elif line and not line.startswith('//') and not line.startswith('*'):
            # Stop when we hit non-comment, non-empty line
            break

    if doxygen_lines:
        combined = '\n'.join(doxygen_lines)
        return parse_doxygen_comment(combined)

    # Look for block comments (/** */ or /*!)
    best_end, best_txt = -1, ""
    for end, txt in comments:
        # Check if it's a Doxygen comment (starts with /** or /*!)
        if end <= pos and (txt.strip().startswith('/**') or txt.strip().startswith('/*!')):
            # Exclude AUTOSAR-generated comments
            if "DO NOT CHANGE THIS COMMENT!" in txt:
                continue

            # Check if there's code (not just whitespace/comments) between comment and function
            between = src[end:pos]
            # Remove whitespace and other comments
            between_clean = re.sub(r'\s+', '', between)
            between_clean = re.sub(r'//.*', '', between_clean)
            between_clean = re.sub(r'/\*.*?\*/', '', between_clean)

            # Only use this comment if there's no significant code between it and the function
            if not between_clean and end > best_end:
                best_end, best_txt = end, txt

    if best_txt:
        return parse_doxygen_comment(best_txt)

    return ""

def parse_macros(src: str) -> list[dict]:
    """Extract #define macros from C source code."""
    # Regex for a macro header: #define NAME [optional(param,list)] body-fragment
    HEADER_RE = re.compile(r"^\s*#define\s+(\w+)\s*(\([^)]*\))?\s*(.*)", re.MULTILINE)

    def should_skip(body: str) -> bool:
        """Return True if the macro should be excluded."""
        stripped = body.strip()
        if not stripped:
            return True
        return False

    lines = src.split('\n')
    macros = []

    i = 0
    while i < len(lines):
        m = HEADER_RE.match(lines[i])
        if not m:
            i += 1
            continue

        line_number = i + 1  # Track line number (1-indexed)
        name = m.group(1)              # macro name
        paramlist = m.group(2)         # None if object-like macro
        first = m.group(3).rstrip()    # first chunk of body

        # Gather continuation lines ending in "\"
        body_parts = [first]
        while body_parts and body_parts[-1].endswith("\\"):
            body_parts[-1] = body_parts[-1][:-1].rstrip()   # drop trailing "\"
            i += 1
            if i >= len(lines):
                break
            body_parts.append(lines[i].rstrip())

        body_text = " ".join(part.strip() for part in body_parts).strip()

        # Remove comments from macro value
        body_text = re.sub(r'/\*.*?\*/', '', body_text)  # Remove /* */ comments
        body_text = re.sub(r'//.*', '', body_text).strip()  # Remove // comments and clean up

        if not should_skip(body_text):
            macros.append({
                "name": name,
                "value": body_text,
                "lineNumber": line_number
            })

        i += 1

    return macros

def parse_variables(src: str) -> list[dict]:
    """Extract global and static global variables from C source code."""
    variables = []

    # Remove preprocessor directives and comments to avoid false matches in body detection
    src_clean = re.sub(r'^\s*#.*$', '', src, flags=re.MULTILINE)
    src_clean = re.sub(r'/\*[\s\S]*?\*/|//.*', '', src_clean)

    # First, identify all function bodies to exclude local variables
    function_bodies = []

    # Find all functions (including FUNC macros, static, and global functions)
    func_patterns = [
        re.compile(r'FUNC\s*\([^)]*\)\s*[A-Za-z_]\w*\s*\([^)]*\)\s*\{', re.MULTILINE),
        re.compile(r'(?:static\s+)?(?:inline\s+)?[A-Za-z_]\w*(?:\s*\*+)?\s+[A-Za-z_]\w*\s*\([^)]*\)\s*\{', re.MULTILINE)
    ]

    for pattern in func_patterns:
        for match in pattern.finditer(src_clean):
            # Find the complete function body
            start_pos = match.end() - 1  # position of opening '{'
            brace_count = 1
            pos = start_pos + 1

            while pos < len(src_clean) and brace_count > 0:
                if src_clean[pos] == '{':
                    brace_count += 1
                elif src_clean[pos] == '}':
                    brace_count -= 1
                pos += 1

            if brace_count == 0:
                function_bodies.append((start_pos, pos))
            else:
                 # If we reached end of file with unclosed braces, don't include this as a function body
                 # This prevents the entire end of file from being marked as "inside function"
                 pass

    def is_inside_function(position):
        """Check if a position is inside any function body."""
        for start, end in function_bodies:
            if start <= position <= end:
                return True
        return False

    # Second, identify all struct/union definition blocks to exclude member variables
    struct_definitions = []

    # Find all struct/union definitions (including typedef'd ones)
    # Matches: struct/union Name { ... } or typedef struct/union { ... } Name_t;
    struct_pattern = re.compile(
        r'(?:typedef\s+)?(?:struct|union)\s*(?:[A-Za-z_]\w*)?\s*\{',
        re.MULTILINE
    )

    for match in struct_pattern.finditer(src_clean):
        # Find the complete struct/union definition body
        start_pos = match.end() - 1  # position of opening '{'
        brace_count = 1
        pos = start_pos + 1

        while pos < len(src_clean) and brace_count > 0:
            if src_clean[pos] == '{':
                brace_count += 1
            elif src_clean[pos] == '}':
                brace_count -= 1
            pos += 1

        if brace_count == 0:
            struct_definitions.append((start_pos, pos))

    def is_inside_struct_definition(position):
        """Check if a position is inside any struct/union definition."""
        for start, end in struct_definitions:
            if start <= position <= end:
                return True
        return False

    def find_in_original_src(pattern, var_name, data_type):
        """
        Find variable declaration in original src to get accurate line number.
        Search for the variable name and data type combination.
        """
        # Create a simpler pattern to find this specific variable in original source
        search_pattern = rf'\b{re.escape(data_type)}\s+{re.escape(var_name)}\b'
        for match in re.finditer(search_pattern, src):
            # Return the position in original source
            return match.start()
        # Fallback: just search for variable name
        search_pattern = rf'\b{re.escape(var_name)}\s*[=;\[]'
        for match in re.finditer(search_pattern, src):
            return match.start()
        return None

    # Debug: Print function bodies for troubleshooting
    #print(f"Function bodies detected: {len(function_bodies)} functions")
    #for start, end in function_bodies:
    #    print(f"  Function body: chars {start}-{end}")

    # Pattern for extern variable declarations
    extern_var_pattern = re.compile(r'''
        ^[ \t]*                                    # start of line, optional whitespace
        (extern\s+)                                # extern keyword
        ([A-Za-z_]\w*(?:\s*\*+)?)                 # data type (with optional pointers)
        \s+                                        # whitespace
        ([A-Za-z_]\w*)                            # variable name
        \s*;(?=\s|$)                                       # semicolon (extern vars don't have initialization)
        ''', re.MULTILINE | re.VERBOSE)
    
    # Pattern for static/regular variable declarations
    static_var_pattern = re.compile(r'''
        ^[ \t]*                                    # start of line, optional whitespace
        (static\s+)?                               # optional 'static' keyword
        ([A-Za-z_]\w*(?:\s*\*+)?)                 # data type (with optional pointers)
        \s+                                        # whitespace
        ([A-Za-z_]\w*)                            # variable name
        (?:\s*=\s*([^;]+))?                       # optional initialization
        \s*;(?=\s|$)                                       # semicolon
        ''', re.MULTILINE | re.VERBOSE)
    
    # Array pattern: type name[size] = {...};
    static_array_pattern = re.compile(r'''
        ^[ \t]*                                    # start of line, optional whitespace
        (static\s+)?                               # optional 'static' keyword
        ([A-Za-z_]\w*(?:\s*\*+)?)                 # data type
        \s+                                        # whitespace
        ([A-Za-z_]\w*)                            # variable name
        \s*\[([^\]]*)\]                           # array brackets with size
        (?:\s*=\s*\{([^}]*)\})?                   # optional array initialization
        \s*;(?=\s|$)                                       # semicolon
        ''', re.MULTILINE | re.VERBOSE)
    
    # Keywords to exclude as variable names (control flow, etc.)
    exclude_var_names = {
        'if', 'for', 'while', 'switch', 'do', 'else', 'case', 'return',
        'goto', 'break', 'continue', 'sizeof'
    }

    # Keywords to exclude as data types (storage classes, type qualifiers that shouldn't be standalone)
    exclude_data_types = {
        'typedef', 'struct', 'union', 'enum', 'extern', 'register',
        'auto', 'volatile', 'const', 'inline'
    }
    
    # Find extern variable declarations
    for match in extern_var_pattern.finditer(src_clean):
        if is_inside_function(match.start()):
            continue

        # Skip if inside struct/union definition
        if is_inside_struct_definition(match.start()):
            continue

        extern_kw = match.group(1)
        data_type = match.group(2).strip()
        var_name = match.group(3)

        if extern_kw and var_name not in exclude_var_names and data_type.lower() not in exclude_data_types:
            # Find position in original source for accurate line number
            src_pos = find_in_original_src(match, var_name, data_type)
            if src_pos is not None:
                line_number = get_line_number(src, src_pos)
                variables.append({
                    "name": var_name,
                    "dataType": data_type,
                    "initialValue": "",
                    "scope": "Extern",
                    "lineNumber": line_number
                })

    # Find static/regular variable declarations
    for match in static_var_pattern.finditer(src_clean):
        static_kw = match.group(1)
        data_type = match.group(2).strip()
        var_name = match.group(3)
        init_value = match.group(4).strip() if match.group(4) else ""

        if is_inside_function(match.start()):
            continue

        # Skip if inside struct/union definition
        if is_inside_struct_definition(match.start()):
            continue

        if var_name in exclude_var_names or data_type.lower() in exclude_data_types:
            continue

        # Check if this looks like a function (has parentheses immediately after the name, before semicolon)
        # The match already ends at semicolon, so we don't need this check anymore
        # (it was causing false positives with macros on subsequent lines)

        # Skip if it looks like a function pointer or typedef
        full_match = match.group(0)
        if '(*' in full_match or 'typedef' in full_match:
            continue

        scope = "Static Global" if static_kw else "Global"

        # Find position in original source for accurate line number
        src_pos = find_in_original_src(match, var_name, data_type)
        if src_pos is not None:
            line_number = get_line_number(src, src_pos)
            variables.append({
                "name": var_name,
                "dataType": data_type,
                "initialValue": init_value,
                "scope": scope,
                "lineNumber": line_number
            })

    # Find array declarations
    for match in static_array_pattern.finditer(src_clean):
        if is_inside_function(match.start()):
            continue

        # Skip if inside struct/union definition
        if is_inside_struct_definition(match.start()):
            continue

        static_kw = match.group(1)
        data_type = match.group(2).strip()
        var_name = match.group(3)
        array_size = match.group(4).strip() if match.group(4) else ""
        init_value = match.group(5).strip() if match.group(5) else ""

        if var_name in exclude_var_names or data_type.lower() in exclude_data_types:
            continue

        scope = "Static Global" if static_kw else "Global"

        # Find position in original source for accurate line number
        src_pos = find_in_original_src(match, var_name, data_type)
        if src_pos is not None:
            line_number = get_line_number(src, src_pos)
            full_type = f"{data_type}[{array_size}]"

            variables.append({
                "name": var_name,
                "dataType": full_type,
                "initialValue": init_value,
                "scope": scope,
                "lineNumber": line_number
            })

    # Post-processing: Clean up any remaining preprocessor keywords from datatypes
    preprocessor_keywords = {
        'endif', 'if', 'ifdef', 'ifndef', 'else', 'elif', 'define', 'include', 
        'undef', 'pragma', 'warning', 'error', 'line'
    }
    
    cleaned_variables = []
    for var in variables:
        # Check if dataType contains any preprocessor keywords
        datatype_words = var['dataType'].lower().split()
        if not any(word in preprocessor_keywords for word in datatype_words):
            cleaned_variables.append(var)
        # Optionally, I can clean individual words instead of removing entire variable:
        # cleaned_datatype = ' '.join(word for word in var['dataType'].split() 
        #                            if word.lower() not in preprocessor_keywords)
        # if cleaned_datatype.strip():
        #     var['dataType'] = cleaned_datatype.strip()
        #     cleaned_variables.append(var)
    
    return cleaned_variables

def get_line_number(src: str, pos: int) -> int:
    """Convert string position to line number (1-indexed)."""
    return src[:pos].count('\n') + 1

def parse_file(src: str, cancel_token: CancellationToken = None) -> tuple[list, list, list]:
    """Parse file and return (functions, macros, variables)."""
    functions = []
    reserved = {"if","for","while","switch","do","else","case","sizeof","abs","return", "endif"}
    exclude_invoked = {
        "VStdLib_MemCpy","VStdLib_MemSet","VStdLib_MemCmp",
        "memcmp","memcpy","memset","sizeof","abs","return"
    }

    # Check for cancellation
    if cancel_token and cancel_token.is_cancelled():
        return [], [], []

    comments = [(m.end(), m.group(0)) for m in re.finditer(r"/\*[\s\S]*?\*/", src)]

    runnable_rx = re.compile(
        r'^[ \t]*FUNC\s*\(\s*([A-Za-z_]\w*)\s*,\s*[A-Za-z_]\w*\s*\)\s*([A-Za-z_]\w*)\s*\(',
        re.MULTILINE
    )
    static_rx = re.compile(r'''
        ^[ \t]*static\s+(?:inline\s+)?                  
        (?:
          FUNC\s*\(\s*([A-Za-z_]\w*)\s*,\s*[A-Za-z_]\w*\s*\)  
        |
          ([A-Za-z_]\w*)                                
        )
        \s+([A-Za-z_]\w*)\s*\(                          
        ''', re.MULTILINE|re.VERBOSE)
    global_rx = re.compile(r'''
        ^[ \t]*(?!static\b)(?!FUNC\b)                   
        ([A-Za-z_]\w*(?:\s*\*+)?)\s+                    
        (?!(?:if|for|while|switch|do|else|case)\b)      
        ([A-Za-z_]\w*)\s*\(                             
        ''', re.MULTILINE|re.VERBOSE)
    sig_rx = re.compile(r'''
        ^[ \t]*                                        
        (?:static\s+)?(?:inline\s+)?                   
        (?:FUNC\([^)]*\)\s*)?                          
        (?P<ret>[\w\*\s]+?)\s+                         
        (?P<name>[A-Za-z_]\w*)\s*                      
        \((?P<params>[^)]*)\)                          
        ''', re.MULTILINE|re.VERBOSE)

    def extract(m, fnType):
        if fnType == "Static":
            retType = m.group(1) or m.group(2)
            name    = m.group(3)
        else:
            retType, name = m.group(1), m.group(2)

        # parameters
        L = len(src)
        depth, i = 1, m.end()
        while i < L and depth:
            depth += src[i] == "("
            depth -= src[i] == ")"
            i += 1
        raw_params = src[m.end():i-1].strip()

        # syntax
        snippet = src[m.start():]
        sig_m = sig_rx.match(snippet)
        if sig_m:
            syntax = f"{sig_m.group('ret').strip()} {sig_m.group('name')}({sig_m.group('params').strip()})"
        else:
            syntax = f"{retType} {name}({raw_params})"

        # skip prototypes
        pos = i
        while pos < L and (src[pos].isspace() or src.startswith("/*", pos)):
            pos = src.find("*/", pos) + 2 if src.startswith("/*", pos) else pos+1
        if pos >= L or src[pos] != "{":
            return

        # body
        brace_idx, depth, j = pos, 1, pos+1
        while j < L and depth:
            depth += src[j] == "{"
            depth -= src[j] == "}"
            j += 1
        body = src[brace_idx+1:j-1]
        code = re.sub(r'/\*[\s\S]*?\*/|//.*', '', body)

        # trigger
        if fnType == "Runnable":
            cm = get_trigger_comment(comments, m.start())
            trigs = re.findall(r"-\s*triggered\s+(?:on|by)\s+([^\n\r]+)", cm, re.IGNORECASE)
            trigger = "; ".join(t.strip() for t in trigs)
        else:
            trigger = ""

        # params names and types
        parts = [p.strip() for p in re.split(r",(?![^(]*\))", raw_params) if p.strip()]
        names = []
        for p in parts:
            # Match parameter name, handling arrays: uint8 arr[] or uint8 arr[10]
            # Also handle pointers: uint8* ptr or uint8 *ptr
            mm = re.search(r"\b([A-Za-z_]\w*)\s*(?:\[[^\]]*\])?\s*$", p)
            names.append(mm.group(1) if mm else p)

        # IN/OUT classification with enhanced pointer analysis
        dirs = classify_params(body, names, parts)

        # AUTOSAR macro overrides
        for orig, nm in zip(parts, names):
            # P2CONST = const pointer, should be IN
            if orig.startswith("P2CONST("):
                dirs[nm] = "IN"
            # P2VAR = variable pointer, can be written to, should be OUT
            elif orig.startswith("P2VAR("):
                dirs[nm] = "OUT"

        inP  = [p for p in names if dirs[p] in ("IN","INOUT")]
        outP = [p for p in names if dirs[p] in ("OUT","INOUT")]

        # RTE APIs
        inputs  = sorted({x.split("(")[0] for x in re.findall(
                     r"\bRte_(?:Read|DRead|IRead|Receive|IReadRef|IrvRead|IsUpdated|Mode_)[\w_]*\s*\(",
                     body)})
        outputs = sorted({x.split("(")[0] for x in re.findall(
                     r"\bRte_(?:Write|IrvWrite|IWrite|IWriteRef|Switch)[\w_]*\s*\(",
                     body)})

        # invoked
        calls = {x.split("(")[0] for x in re.findall(r"\bRte_Call_[\w_]+\s*\(", code)}
        plain = re.findall(r"\b([A-Za-z_]\w*)\s*\(", code)
        locals_ = {
            c for c in plain
            if c not in reserved
            and c not in exclude_invoked
            and not c.startswith("Rte_")
            and c != name
        }
        invoked = sorted(c for c in calls|locals_ if not re.fullmatch(r"[A-Z][A-Z0-9_]*", c))

        # used types
        used = sorted({
            t for t in re.findall(r"\b([A-Za-z_]\w*)\s+[A-Za-z_]\w*\s*(?:[=;])", body)
            if t.lower() not in reserved
        })

        # placeholders for GUI fields
        sync_async = ""
        reentrancy = ""

        # Extract Doxygen description from comments above function
        description = get_doxygen_comment(src, comments, m.start())

        # Calculate line number
        line_number = get_line_number(src, m.start())

        functions.append({
            "name":       name,
            "syntax":     syntax,
            "ret":        retType,
            "inParams":   inP,
            "outParams":  outP,
            "fnType":     fnType,
            "trigger":    trigger,
            "inputs":     inputs,
            "outputs":    outputs,
            "invoked":    invoked,
            "used":       used,
            "Sync_Async": sync_async,
            "Reentrancy": reentrancy,
            "description": description,
            "lineNumber": line_number
        })

    for m in runnable_rx.finditer(src):
        if cancel_token and cancel_token.is_cancelled():
            return [], [], []
        extract(m, "Runnable")

    for m in static_rx.finditer(src):
        if cancel_token and cancel_token.is_cancelled():
            return [], [], []
        extract(m, "Static")

    for m in global_rx.finditer(src):
        if cancel_token and cancel_token.is_cancelled():
            return [], [], []
        extract(m, "Global")

    # Check cancellation before parsing macros and variables
    if cancel_token and cancel_token.is_cancelled():
        return [], [], []

    # Parse macros and variables
    macros = parse_macros(src)
    variables = parse_variables(src)

    return functions, macros, variables

def show_gui_old():
    function_fields = [
      "Line Number", "Name", "Description", "Syntax", "Triggers", "In-Parameters", "Out-Parameters",
      "Return Value", "Function Type", "Inputs", "Outputs",
      "Invoked Operations", "Used Data Types", "Sync/Async", "Reentrancy"
    ]

    macro_fields = [
        "Line Number", "Name", "Value"
    ]

    variable_fields = [
        "Line Number", "Name", "Data Type", "Initial Value", "Scope"
    ]

    formats = ["Excel", "Word", "MD"]

    root = tk.Tk()
    root.title("Documentation Slayer")
    root.configure(bg="#ececec")
    if getattr(sys, "frozen", False):
        # running from PyInstaller bundle
        base_path = Path(sys._MEIPASS)
    else:
        # running as a normal script
        base_path = Path(__file__).parent
    icon_path = base_path / "DocSlayerLogo.ico"
    try:
        root.iconbitmap(str(icon_path))
    except Exception:
        pass  # skip if missing or invalid
    root.geometry("900x600")
    root.minsize(900, 600)
    root.maxsize(900, 600)
    root.resizable(False, False) # no resizing

    root.update_idletasks()         # ensure winfo_width/height are accurate
    w = root.winfo_width()
    h = root.winfo_height()
    sw = root.winfo_screenwidth()
    sh = root.winfo_screenheight()
    x = (sw - w) // 2
    y = (sh - h) // 2
    root.geometry(f"{w}x{h}+{x}+{y}")

    # Create notebook for tabs
    notebook = ttk.Notebook(root)
    notebook.pack(fill="both", expand=True, padx=10, pady=10)

    # Functions tab
    functions_frame = ttk.Frame(notebook)
    notebook.add(functions_frame, text="Functions")

    function_vars = {f: tk.BooleanVar(value=(f != "Line Number")) for f in function_fields}

    ttk.Label(functions_frame, text="Select function fields:").grid(row=0, column=0, sticky="w", padx=10, pady=5)

    # Select All / Deselect All buttons
    button_frame_func = ttk.Frame(functions_frame)
    button_frame_func.grid(row=0, column=1, columnspan=2, sticky="e", padx=10, pady=5)

    def select_all_functions():
        for var in function_vars.values():
            var.set(True)

    def deselect_all_functions():
        for var in function_vars.values():
            var.set(False)

    ttk.Button(button_frame_func, text="Select All", command=select_all_functions).pack(side=tk.LEFT, padx=5)
    ttk.Button(button_frame_func, text="Deselect All", command=deselect_all_functions).pack(side=tk.LEFT, padx=5)

    for i, f in enumerate(function_fields):
        chk = ttk.Checkbutton(functions_frame, text=f, variable=function_vars[f])
        chk.grid(row=(i//3)+1, column=i%3, sticky="w", padx=15, pady=2)

    # Macros tab
    macros_frame = ttk.Frame(notebook)
    notebook.add(macros_frame, text="Macros")

    macro_vars = {f: tk.BooleanVar(value=(f != "Line Number")) for f in macro_fields}

    ttk.Label(macros_frame, text="Select macro fields:").grid(row=0, column=0, sticky="w", padx=10, pady=5)

    # Select All / Deselect All buttons
    button_frame_macro = ttk.Frame(macros_frame)
    button_frame_macro.grid(row=0, column=1, columnspan=2, sticky="e", padx=10, pady=5)

    def select_all_macros():
        for var in macro_vars.values():
            var.set(True)

    def deselect_all_macros():
        for var in macro_vars.values():
            var.set(False)

    ttk.Button(button_frame_macro, text="Select All", command=select_all_macros).pack(side=tk.LEFT, padx=5)
    ttk.Button(button_frame_macro, text="Deselect All", command=deselect_all_macros).pack(side=tk.LEFT, padx=5)

    for i, f in enumerate(macro_fields):
        chk = ttk.Checkbutton(macros_frame, text=f, variable=macro_vars[f])
        chk.grid(row=(i//3)+1, column=i%3, sticky="w", padx=15, pady=2)

    # Variables tab
    variables_frame = ttk.Frame(notebook)
    notebook.add(variables_frame, text="Variables")

    variable_vars = {f: tk.BooleanVar(value=(f != "Line Number")) for f in variable_fields}

    ttk.Label(variables_frame, text="Select variable fields:").grid(row=0, column=0, sticky="w", padx=10, pady=5)

    # Select All / Deselect All buttons
    button_frame_var = ttk.Frame(variables_frame)
    button_frame_var.grid(row=0, column=1, columnspan=2, sticky="e", padx=10, pady=5)

    def select_all_variables():
        for var in variable_vars.values():
            var.set(True)

    def deselect_all_variables():
        for var in variable_vars.values():
            var.set(False)

    ttk.Button(button_frame_var, text="Select All", command=select_all_variables).pack(side=tk.LEFT, padx=5)
    ttk.Button(button_frame_var, text="Deselect All", command=deselect_all_variables).pack(side=tk.LEFT, padx=5)

    for i, f in enumerate(variable_fields):
        chk = ttk.Checkbutton(variables_frame, text=f, variable=variable_vars[f])
        chk.grid(row=(i//3)+1, column=i%3, sticky="w", padx=15, pady=2)

    # Activity Diagram tab (password protected)
    activity_frame = ttk.Frame(notebook)
    
    def on_activity_tab_selected(event):
        selected_tab = event.widget.tab('current')['text']
        if selected_tab == "Activity Diagram" and not password_manager.is_authenticated:
            if not ask_password(root):
                # If password wrong or cancelled, go back to Variables tab
                notebook.select(2)  # Variables tab index
                return
    
    notebook.bind("<<NotebookTabChanged>>", on_activity_tab_selected)
    notebook.add(activity_frame, text="Activity Diagram")
    
    # Activity Diagram content with button instead of checkbox
    ttk.Label(activity_frame, text="Activity Diagram Generator", font=("Arial", 14, "bold")).pack(pady=20)
    ttk.Label(activity_frame, text="Click the button below to generate activity diagrams from your C source file.").pack(pady=10)
    
    
    def generate_activity_diagram():
        if not password_manager.is_authenticated:
            if not ask_password(root):
                return
        
        # # Ask for C file if not already selected
        # if not selected_file["path"]:
        #     cfile = filedialog.askopenfilename(
        #         title="Select C source file for activity diagram",
        #         filetypes=[("C files","*.c"), ("All files","*.*")]
        #     )
        #     if not cfile:
        #         return
        #     selected_file["path"] = cfile
        
        # # Get output directory
        # outdir = Path(save_dir.get())
        # stem = Path(selected_file["path"]).stem
        
        # Execute the external CodeSmasher.exe
        try:
            if getattr(sys, "frozen", False):
                base_path = Path(sys._MEIPASS)
            else:
                base_path = Path(__file__).parent
            
            exe_path = base_path / "CodeSmasher.exe"
            if exe_path.exists():
                # Run CodeSmasher.exe 
                cmd = [str(exe_path)]
                result = subprocess.run(cmd, capture_output=True, text=True)
                if result.returncode == 0:
                    messagebox.showinfo("Success", "Activity diagrams got Slayed (generated) successfully!")
                else:
                    messagebox.showerror("Error", f"Failed to generate activity diagrams:\n{result.stderr}")
            else:
                messagebox.showerror("Error", "CodeSmasher.exe not found! Please ensure it's in the same directory as this script.")
        except Exception as e:
            messagebox.showerror("Error", f"Error running CodeSmasher.exe: {str(e)}")
    
    generate_btn = ttk.Button(activity_frame, text="Generate Activity Diagram", command=generate_activity_diagram, style="Accent.TButton")
    generate_btn.pack(pady=30)
    
    # Add style for accent button
    style = ttk.Style()
    style.configure("Accent.TButton", font=("Arial", 11, "bold"))

    # Settings frame (at bottom)
    settings_frame = ttk.Frame(root)
    settings_frame.pack(fill="x", padx=10, pady=(0, 10))

    format_vars = {f: tk.BooleanVar(value=True) for f in formats}
    save_dir = tk.StringVar(value=str(Path.cwd()))

    ttk.Label(settings_frame, text="Select formats:").grid(row=0, column=0, sticky="w", pady=5)
    for i, fmt in enumerate(formats):
        chk = ttk.Checkbutton(settings_frame, text=fmt, variable=format_vars[fmt])
        chk.grid(row=0, column=i+1, sticky="w", padx=15, pady=5)

    def choose_dir():
        d = filedialog.askdirectory()
        if d:
            save_dir.set(d)
    ttk.Button(settings_frame, text="Save to…", command=choose_dir).grid(row=1, column=0, sticky="w", pady=5)
    ttk.Label(settings_frame, textvariable=save_dir).grid(row=1, column=1, columnspan=3, sticky="w")

    def on_run():
        # Gather user selections
        sel_function_fields = [f for f, v in function_vars.items() if v.get()]
        sel_macro_fields = [f for f, v in macro_vars.items() if v.get()]
        sel_variable_fields = [f for f, v in variable_vars.items() if v.get()]
        sel_formats = [f for f, v in format_vars.items() if v.get()]
        outdir = Path(save_dir.get())

        # Ask for the C source file
        cfile = filedialog.askopenfilename(
            title="Select C source file",
            filetypes=[("C files","*.c"), ("All files","*.*")]
        )
        if not cfile:
            return

        # Create progress dialog and cancellation token
        progress = ProgressDialog(root, "Processing C File")
        cancel_token = CancellationToken()
        result = {"success": False, "error": None, "functions": [], "macros": [], "variables": [], "stem": ""}

        def process_in_thread():
            try:
                # Read file
                progress.update_text("Reading file...")
                src = open(cfile, encoding="utf-8").read()

                # Parse the file
                progress.update_text(f"Parsing file ({len(src):,} bytes)...")
                functions, macros, variables = parse_file(src, cancel_token)

                if cancel_token.is_cancelled() or progress.cancelled:
                    result["error"] = "cancelled"
                    return

                result["functions"] = functions
                result["macros"] = macros
                result["variables"] = variables
                result["stem"] = Path(cfile).stem

                # Export files
                xlsx_path = outdir / f"{result['stem']}.xlsx"

                if "Excel" in sel_formats:
                    if cancel_token.is_cancelled() or progress.cancelled:
                        result["error"] = "cancelled"
                        return
                    progress.update_text("Exporting to Excel...")
                    write_excel(str(xlsx_path), functions, macros, variables,
                               sel_function_fields, sel_macro_fields, sel_variable_fields)

                if "Word" in sel_formats:
                    if cancel_token.is_cancelled() or progress.cancelled:
                        result["error"] = "cancelled"
                        return
                    progress.update_text("Exporting to Word...")
                    write_docx(str(xlsx_path), result['stem'], functions, macros, variables,
                              sel_function_fields, sel_macro_fields, sel_variable_fields)

                if "MD" in sel_formats:
                    if cancel_token.is_cancelled() or progress.cancelled:
                        result["error"] = "cancelled"
                        return
                    progress.update_text("Exporting to Markdown...")
                    md_path = outdir / f"{result['stem']}.md"
                    write_markdown(str(md_path), functions, macros, variables,
                                  sel_function_fields, sel_macro_fields, sel_variable_fields)

                result["success"] = True

            except Exception as e:
                result["error"] = str(e)

        # Start processing in a separate thread
        thread = threading.Thread(target=process_in_thread, daemon=True)
        thread.start()

        # Wait for thread to complete
        while thread.is_alive():
            root.update()
            if progress.cancelled:
                cancel_token.cancel()
                break
            thread.join(timeout=0.1)

        # Close progress dialog
        progress.close()

        # Handle results
        if result["error"] == "cancelled":
            messagebox.showwarning("Cancelled", "Operation was cancelled by user.")
            return
        elif result["error"]:
            messagebox.showerror("Error", f"An error occurred:\n{result['error']}")
            return
        elif result["success"]:
            # Open generated files
            xlsx_path = outdir / f"{result['stem']}.xlsx"
            if "Excel" in sel_formats:
                open_file(str(xlsx_path))
            if "Word" in sel_formats:
                open_file(str(xlsx_path.with_suffix('.docx')))
            if "MD" in sel_formats:
                md_path = outdir / f"{result['stem']}.md"
                open_file(str(md_path))

            # Show success message
            messagebox.showinfo("Success", f"Documentation got Slayed (generated) successfully!\n\nFile: {result['stem']}\nFormats: {', '.join(sel_formats)}\n")

    ttk.Button(settings_frame, text="Run", command=on_run).grid(row=2, column=0, pady=15)
    root.mainloop()


def show_gui():
    """Launch the PyQt6 GUI (Ultra Modern Edition)"""
    from qt_gui_modern import DocumentationSlayerModernGUI
    from splash_screen import ModernSplashScreen

    app = QApplication(sys.argv)

    # Apply modern stylesheet
    app.setStyle('Fusion')

    # Show splash screen
    if getattr(sys, "frozen", False):
        base_path = Path(sys._MEIPASS)
    else:
        base_path = Path(__file__).parent

    logo_path = base_path / "DocSlayerLogo.ico"
    splash = ModernSplashScreen(str(logo_path) if logo_path.exists() else None)
    splash.show()
    app.processEvents()

    # Create main window
    window = DocumentationSlayerModernGUI(
        password_manager=password_manager,
        parse_file_func=parse_file,
        write_excel_func=write_excel,
        write_docx_func=write_docx,
        write_markdown_func=write_markdown,
        open_file_func=open_file
    )

    # Show main window after splash
    def show_main_window():
        window.show()
        splash.close()

    # Delay showing main window
    QTimer.singleShot(3000, show_main_window)

    sys.exit(app.exec())


def log_verbose(message, verbose=False):
    """Print verbose logging messages"""
    if verbose:
        print(f"[INFO] {message}", file=sys.stderr)


def process_file_cli(file_path, parse_types, output_path, output_formats, verbose=False, cancel_token=None):
    """Process a single file in CLI mode"""
    log_verbose(f"Processing file: {file_path}", verbose)

    try:
        with open(file_path, encoding="utf-8") as f:
            src = f.read()
    except Exception as e:
        print(f"❌ Error reading {file_path}: {e}", file=sys.stderr)
        return False

    # Show file size for progress estimation
    file_size = len(src)
    if verbose:
        print(f"[INFO] File size: {file_size:,} bytes ({file_size / 1024:.1f} KB)", file=sys.stderr)

    try:
        if verbose and TQDM_AVAILABLE:
            print("[INFO] Parsing file...", file=sys.stderr)
        functions, macros, variables = parse_file(src, cancel_token)

        # Check if cancelled
        if cancel_token and cancel_token.is_cancelled():
            print("⚠️  Operation cancelled by user", file=sys.stderr)
            return False
    except Exception as e:
        print(f"❌ Error parsing {file_path}: {e}", file=sys.stderr)
        return False

    # Filter based on parse_types
    if 'all' not in parse_types:
        if 'functions' not in parse_types:
            functions = []
        if 'macros' not in parse_types:
            macros = []
        if 'variables' not in parse_types:
            variables = []

    log_verbose(f"Parsed {len(functions)} functions, {len(macros)} macros, {len(variables)} variables", verbose)

    # Common field selections
    function_fields = ["Line Number","Name","Description","Syntax","Triggers","In-Parameters","Out-Parameters",
                      "Return Value","Function Type","Inputs","Outputs",
                      "Invoked Operations","Used Data Types","Sync/Async","Reentrancy"]
    macro_fields = ["Line Number", "Name", "Value"]
    variable_fields = ["Line Number", "Name", "Data Type", "Initial Value", "Scope"]

    # Export to each requested format
    all_success = True

    # Use tqdm for progress if available and verbose
    format_iterator = tqdm(output_formats, desc="Exporting", disable=not (verbose and TQDM_AVAILABLE)) if TQDM_AVAILABLE else output_formats

    for output_format in format_iterator:
        # Determine output path for this format
        if output_path:
            # If output_path is a directory, create file inside it
            output_path_obj = Path(output_path)
            if output_path_obj.is_dir() or not output_path_obj.suffix:
                output_path_obj.mkdir(parents=True, exist_ok=True)
                file_stem = Path(file_path).stem
                if output_format == 'excel':
                    format_output_path = str(output_path_obj / f"{file_stem}_documentation.xlsx")
                elif output_format == 'markdown':
                    format_output_path = str(output_path_obj / f"{file_stem}_documentation.md")
                elif output_format == 'word':
                    format_output_path = str(output_path_obj / f"{file_stem}_documentation.docx")
                else:
                    format_output_path = str(output_path_obj / f"{file_stem}_documentation.json")
            else:
                # Use the provided path as-is (for single format)
                format_output_path = output_path
        else:
            # No output path specified, save next to input file
            file_stem = Path(file_path).stem
            file_dir = Path(file_path).parent
            if output_format == 'excel':
                format_output_path = str(file_dir / f"{file_stem}_documentation.xlsx")
            elif output_format == 'markdown':
                format_output_path = str(file_dir / f"{file_stem}_documentation.md")
            elif output_format == 'word':
                format_output_path = str(file_dir / f"{file_stem}_documentation.docx")
            else:
                format_output_path = str(file_dir / f"{file_stem}_documentation.json")

        # Export based on format
        try:
            if output_format == 'excel':
                log_verbose(f"Exporting to Excel: {format_output_path}", verbose)
                write_excel(format_output_path, functions, macros, variables,
                    function_fields, macro_fields, variable_fields)

            elif output_format == 'markdown':
                log_verbose(f"Exporting to Markdown: {format_output_path}", verbose)
                write_markdown(format_output_path, functions, macros, variables,
                    function_fields, macro_fields, variable_fields)

            elif output_format == 'word':
                log_verbose(f"Exporting to Word: {format_output_path}", verbose)
                swcName = Path(file_path).stem
                write_docx(format_output_path, swcName, functions, macros, variables,
                    function_fields, macro_fields, variable_fields)

            elif output_format == 'json':
                log_verbose(f"Exporting to JSON: {format_output_path}", verbose)
                output = {
                    "functions": functions,
                    "macros": macros,
                    "variables": variables
                }
                with open(format_output_path, 'w', encoding='utf-8') as f:
                    json.dump(output, f, indent=2)

            log_verbose(f"✓ Successfully exported to {format_output_path}", verbose)

        except Exception as e:
            print(f"❌ Error exporting to {output_format}: {e}", file=sys.stderr)
            all_success = False

    return all_success


def process_directory_cli(dir_path, parse_types, output_path, output_formats, file_pattern, recursive, verbose=False):
    """Process all matching files in a directory"""
    log_verbose(f"Scanning directory: {dir_path} (recursive={recursive})", verbose)

    path = Path(dir_path)
    if recursive:
        pattern_files = list(path.rglob(file_pattern))
    else:
        pattern_files = list(path.glob(file_pattern))

    log_verbose(f"Found {len(pattern_files)} files matching pattern '{file_pattern}'", verbose)

    success_count = 0
    for file_path in pattern_files:
        # Generate unique output path for each file
        if output_path:
            # Create output in specified directory with same structure
            rel_path = file_path.relative_to(path)
            out_dir = Path(output_path)
            out_file = out_dir / rel_path.parent
        else:
            out_file = file_path.parent

        # Ensure output directory exists
        out_file.mkdir(parents=True, exist_ok=True)

        if process_file_cli(str(file_path), parse_types, str(out_file), output_formats, verbose):
            success_count += 1

    log_verbose(f"Processed {success_count}/{len(pattern_files)} files successfully", verbose)
    return success_count == len(pattern_files)


def process_batch_config(config_file, verbose=False):
    """Process multiple files based on JSON config"""
    log_verbose(f"Loading batch config: {config_file}", verbose)

    try:
        with open(config_file, 'r', encoding='utf-8') as f:
            config = json.load(f)
    except Exception as e:
        print(f"❌ Error reading config file {config_file}: {e}", file=sys.stderr)
        return False

    inputs = config.get('inputs', [])
    output = config.get('output', None)
    output_format = config.get('format', 'excel')
    parse_types = config.get('parse', ['all'])
    file_pattern = config.get('file_pattern', '*.c')
    recursive = config.get('recursive', False)

    # Parse formats (support both string and list)
    if isinstance(output_format, str):
        output_formats = [f.strip().lower() for f in output_format.split(',')]
    else:
        output_formats = output_format

    log_verbose(f"Batch config: {len(inputs)} inputs, format={','.join(output_formats)}", verbose)

    all_success = True
    for input_path in inputs:
        input_path_obj = Path(input_path)

        if input_path_obj.is_file():
            if not process_file_cli(input_path, parse_types, output, output_formats, verbose):
                all_success = False
        elif input_path_obj.is_dir():
            if not process_directory_cli(input_path, parse_types, output, output_formats, file_pattern, recursive, verbose):
                all_success = False
        else:
            print(f"❌ Input path does not exist: {input_path}", file=sys.stderr)
            all_success = False

    return all_success


def main():
    ap = argparse.ArgumentParser(
        description="Documentation Slayer - Extract and document C functions, macros, and variables",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Launch GUI (default)
  python parser.py

  # Parse single file to Excel
  python parser.py --input file.c --format excel

  # Parse single file to multiple formats (Excel and Word)
  python parser.py --input file.c --format excel,word --output docs/

  # Parse directory recursively to Markdown
  python parser.py --input src/ --recursive --format markdown --parse functions,macros

  # Use config file for batch processing
  python parser.py --config batch.json --verbose

  # CLI mode with multiple formats and verbose output
  python parser.py --input file.c --no-gui --format excel,word,markdown --output docs/ --verbose
        """
    )

    # Legacy support for positional file argument
    ap.add_argument("file", nargs="?", default=None, help="C source file path (legacy, use --input instead)")

    # Main options
    ap.add_argument("--input", "-i", help="Input file or directory path")
    ap.add_argument("--output", "-o", help="Output file or directory path")
    ap.add_argument("--format", "-f", default='excel',
                    help="Export format(s): excel, word, markdown (comma-separated, default: excel)")
    ap.add_argument("--parse", "-p", default='all',
                    help="What to parse: functions, macros, variables, all (comma-separated, default: all)")
    ap.add_argument("--no-gui", "--nogui", action="store_true",
                    help="Run without GUI (CLI mode only)")
    ap.add_argument("--config", "-c", help="JSON config file for batch processing")
    ap.add_argument("--recursive", "-r", action="store_true",
                    help="Parse directories recursively")
    ap.add_argument("--file-pattern", default="*.c",
                    help="File pattern to match (default: *.c)")
    ap.add_argument("--verbose", "-v", action="store_true",
                    help="Enable verbose logging")

    args = ap.parse_args()

    # Parse types
    parse_types = [t.strip().lower() for t in args.parse.split(',')]

    # Parse formats (comma-separated)
    output_formats = [f.strip().lower() for f in args.format.split(',')]

    # Config file mode
    if args.config:
        success = process_batch_config(args.config, args.verbose)
        sys.exit(0 if success else 1)

    # Determine input (--input takes precedence over positional file argument)
    input_path = args.input or args.file

    # If no input and no --no-gui, launch GUI
    if not input_path and not args.no_gui:
        show_gui()
        sys.exit(0)

    # CLI mode requires input
    if not input_path:
        print("❌ Error: --input required in CLI mode (or use GUI without --no-gui)", file=sys.stderr)
        ap.print_help()
        sys.exit(1)

    # Process input
    input_path_obj = Path(input_path)

    if input_path_obj.is_file():
        success = process_file_cli(input_path, parse_types, args.output, output_formats, args.verbose)
        sys.exit(0 if success else 1)

    elif input_path_obj.is_dir():
        success = process_directory_cli(input_path, parse_types, args.output, output_formats,
                                       args.file_pattern, args.recursive, args.verbose)
        sys.exit(0 if success else 1)

    else:
        print(f"❌ Error: Input path does not exist: {input_path}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()